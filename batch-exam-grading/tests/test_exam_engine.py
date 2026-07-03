import json
import os
import shutil
import tempfile
import unittest
from unittest import mock

from docx import Document
import openpyxl

try:
    from scripts.exam_engine import build_llm_requests, grade_roster, load_exam_config
    from scripts import grade_exam as grade_exam_module
    from scripts import llm_grade as llm_grade_module
except ImportError:
    build_llm_requests = None
    grade_exam_module = None
    grade_roster = None
    load_exam_config = None
    llm_grade_module = None


class ExamEngineTests(unittest.TestCase):
    def setUp(self):
        self.tempdir = tempfile.mkdtemp(prefix="batch-exam-grading-")
        self.answer_dir = os.path.join(self.tempdir, "answers")
        os.makedirs(self.answer_dir, exist_ok=True)
        self.config_path = os.path.join(self.tempdir, "exam_config.json")
        self.roster_path = os.path.join(self.tempdir, "roster.csv")
        self.roster_xlsx_path = os.path.join(self.tempdir, "roster.xlsx")
        self.llm_path = os.path.join(self.tempdir, "llm_grades.jsonl")
        self.llm_cache_path = os.path.join(self.tempdir, "llm_cache.jsonl")
        self.output_path = os.path.join(self.tempdir, "scores.xlsx")
        self.docx_path = os.path.join(self.answer_dir, "2025001_sample.docx")

        with open(self.roster_path, "w", encoding="utf-8-sig", newline="") as f:
            f.write("index,sid,student_name\n")
            f.write("1,2025001,张三\n")
        wb = openpyxl.Workbook()
        ws = wb.active
        ws.append(["序号", "学号", "姓名"])
        ws.append([1, "2025001", "张三"])
        wb.save(self.roster_xlsx_path)

        doc = Document()
        doc.add_paragraph("一、客观题")
        doc.add_paragraph("1、A 2、虚拟机")
        doc.add_paragraph("二、程序分析")
        doc.add_paragraph("1.（5分） 输出结果是 42")
        doc.add_paragraph("三、编程题")
        table = doc.add_table(rows=1, cols=1)
        table.cell(0, 0).text = "public class Demo { public static void main(String[] args){ System.out.println(42); } }"
        doc.save(self.docx_path)

        config = {
            "files": {
                "answer_folder": self.answer_dir,
                "roster_file": self.roster_xlsx_path,
                "output_xlsx": self.output_path,
                "llm_grades_jsonl": self.llm_path,
                "llm_cache_jsonl": self.llm_cache_path
            },
            "roster": {
                "student_id_field": "学号",
                "name_field": "姓名",
                "seq_field": "序号"
            },
            "parsing": {
                "student_id_pattern": r"(\\d{7})",
                "sections": {
                    "choice": {
                        "start_keywords": ["客观题"],
                        "end_keywords": ["程序分析"]
                    },
                    "analysis": {
                        "start_keywords": ["程序分析"],
                        "end_keywords": ["编程题"]
                    }
                },
                "code_table_keywords": ["class", "public", "System.out"],
                "code_min_length": 20
            },
            "llm": {
                "mode": "llm_api"
            },
            "questions": [
                {
                    "id": 1,
                    "type": "choice",
                    "score": 2,
                    "section": "一、客观题",
                    "extract": {"source": "choice_inline"},
                    "grading": {"mode": "exact", "answer": "A"}
                },
                {
                    "id": 2,
                    "type": "fill",
                    "score": 2,
                    "section": "一、客观题",
                    "extract": {"source": "choice_inline", "answer_pattern": r"2、([^\\s]+)"},
                    "grading": {"mode": "contains_any", "answers": ["虚拟机", "JVM"]}
                },
                {
                    "id": 3,
                    "type": "analysis",
                    "score": 5,
                    "section": "二、程序分析",
                    "extract": {"source": "section_question", "question_number": 1},
                    "llm": {
                        "enabled": True,
                        "prompt": "分析题题干",
                        "reference_answer": "42",
                        "rubric": [{"point": "结果正确", "score": 5}]
                    },
                    "fallback_grading": {"mode": "contains_any", "answers": ["42"]}
                },
                {
                    "id": 4,
                    "type": "code",
                    "score": 4,
                    "section": "三、编程题",
                    "extract": {"source": "code_table", "table_index": 0},
                    "llm": {
                        "enabled": True,
                        "prompt": "编程题题干",
                        "reference_answer": "输出 42",
                        "rubric": [{"point": "输出正确", "score": 4}]
                    },
                    "fallback_grading": {"mode": "keyword_points", "points": [{"keywords": ["System.out.println", "42"], "score": 4}]}
                }
            ]
        }
        with open(self.config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)

        request_hashes = {}
        if load_exam_config is not None and build_llm_requests is not None:
            loaded = load_exam_config(self.config_path)
            for item in build_llm_requests(loaded):
                request_hashes[item["question_id"]] = item["request_hash"]

        with open(self.llm_path, "w", encoding="utf-8") as f:
            f.write(json.dumps({
                "student_id": "2025001",
                "question_id": 3,
                "question_type": "analysis",
                "max_score": 5,
                "request_hash": request_hashes.get(3, ""),
                "score": 5,
                "confidence": "high",
                "needs_human_review": False,
                "deductions": [],
                "evidence": ["输出结果为 42"],
                "filename": "2025001_sample.docx"
            }, ensure_ascii=False) + "\n")
            f.write(json.dumps({
                "student_id": "2025001",
                "question_id": 4,
                "question_type": "code",
                "max_score": 4,
                "request_hash": request_hashes.get(4, ""),
                "score": 4,
                "confidence": "high",
                "needs_human_review": False,
                "deductions": [],
                "evidence": ["代码输出 42"],
                "filename": "2025001_sample.docx"
            }, ensure_ascii=False) + "\n")

    def tearDown(self):
        try:
            shutil.rmtree(self.tempdir)
        except PermissionError:
            pass

    def test_grade_exam_uses_question_list_and_roster_field_mapping(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(grade_roster, "grade_roster 未实现")
        config = load_exam_config(self.config_path)
        result = grade_roster(config)
        self.assertEqual(result["scores"]["2025001"]["q1"], 2)
        self.assertEqual(result["scores"]["2025001"]["q2"], 2)
        self.assertEqual(result["scores"]["2025001"]["q3"], 5)
        self.assertEqual(result["scores"]["2025001"]["q4"], 4)
        self.assertEqual(result["roster"][0]["name"], "张三")

    def test_collect_llm_requests_reads_subjective_questions_from_config(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(build_llm_requests, "build_llm_requests 未实现")
        config = load_exam_config(self.config_path)
        requests = build_llm_requests(config)
        self.assertEqual([item["question_id"] for item in requests], [3, 4])
        self.assertEqual(requests[0]["max_score"], 5)
        self.assertEqual(requests[1]["max_score"], 4)
        self.assertIn("request_hash", requests[0])

    def test_grade_exam_requires_llm_scores_for_subjective_questions(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(grade_roster, "grade_roster 未实现")
        os.remove(self.llm_path)
        config = load_exam_config(self.config_path)
        with self.assertRaisesRegex(ValueError, "缺少主观题模型评分"):
            grade_roster(config)

    def test_grade_with_openai_reuses_cached_result_for_same_request_hash(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(build_llm_requests, "build_llm_requests 未实现")
        self.assertIsNotNone(llm_grade_module, "llm_grade 模块未实现")
        config = load_exam_config(self.config_path)
        requests = build_llm_requests(config)
        cached_rows = [{
            "student_id": "2025001",
            "anonymous_student_id": "sample",
            "filename": "2025001_sample.docx",
            "question_id": requests[0]["question_id"],
            "question_type": requests[0]["question_type"],
            "max_score": requests[0]["max_score"],
            "request_hash": requests[0]["request_hash"],
            "score": 5,
            "confidence": "high",
            "needs_human_review": False,
            "deductions": [],
            "evidence": ["cached"]
        }]
        with mock.patch.object(llm_grade_module, "call_llm_api", side_effect=AssertionError("不应调用模型")):
            rows = llm_grade_module.grade_with_llm_api(
                requests[:1],
                {"provider": "openai", "model": "dummy-model"},
                review_policy={},
                cached_results={requests[0]["request_hash"]: cached_rows[0]}
            )
        self.assertEqual(rows[0]["score"], 5)
        self.assertEqual(rows[0]["evidence"], ["cached"])

    def test_grade_exam_supports_csv_roster_input(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(grade_roster, "grade_roster 未实现")
        config = load_exam_config(self.config_path)
        config["files"]["roster_file"] = self.roster_path
        config["roster"] = {
            "student_id_field": "sid",
            "name_field": "student_name",
            "seq_field": "index"
        }
        result = grade_roster(config)
        self.assertEqual(result["roster"][0]["student_id"], "2025001")
        self.assertEqual(result["roster"][0]["name"], "张三")

    def test_grade_exam_supports_legacy_roster_csv_field(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(grade_roster, "grade_roster 未实现")
        config = load_exam_config(self.config_path)
        config["files"].pop("roster_file", None)
        config["files"]["roster_csv"] = self.roster_xlsx_path
        with open(self.config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        reloaded = load_exam_config(self.config_path)
        result = grade_roster(reloaded)
        self.assertEqual(reloaded["files"]["roster_file"], self.roster_xlsx_path)
        self.assertEqual(result["roster"][0]["student_id"], "2025001")
        self.assertEqual(result["roster"][0]["name"], "张三")

    def test_roster_file_takes_precedence_over_legacy_roster_csv(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(grade_roster, "grade_roster 未实现")
        config = load_exam_config(self.config_path)
        config["files"]["roster_file"] = self.roster_xlsx_path
        config["files"]["roster_csv"] = self.roster_path
        with open(self.config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        reloaded = load_exam_config(self.config_path)
        result = grade_roster(reloaded)
        self.assertEqual(reloaded["files"]["roster_file"], self.roster_xlsx_path)
        self.assertEqual(result["roster"][0]["student_id"], "2025001")
        self.assertEqual(result["roster"][0]["name"], "张三")

    def test_section_regex_extracts_answer_within_target_section(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        doc = Document()
        doc.add_paragraph("一、客观题")
        doc.add_paragraph("1、A")
        doc.add_paragraph("二、填空题")
        doc.add_paragraph("1、JVM")
        doc.add_paragraph("2、虚拟机")
        doc.add_paragraph("三、程序分析")
        doc.save(self.docx_path)
        config = load_exam_config(self.config_path)
        config["questions"][1]["extract"] = {
            "source": "section_regex",
            "section_name": "fill",
            "pattern": "2[、.．]\\s*(.+)"
        }
        config["parsing"]["sections"]["fill"] = {
            "start_keywords": ["二、填空题"],
            "end_keywords": ["三、程序分析"]
        }
        config.setdefault("llm", {})["require_for_subjective"] = False
        config["questions"] = config["questions"][:2]
        result = grade_roster(config)
        self.assertEqual(result["scores"]["2025001"]["q2"], 2)

    def test_write_excel_aggregates_multiple_questions_by_same_section(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(grade_roster, "grade_roster 未实现")
        try:
            from scripts.exam_engine import write_excel
        except ImportError:
            from exam_engine import write_excel
        config = load_exam_config(self.config_path)
        result = grade_roster(config)
        output_path = write_excel(config, result)
        wb = openpyxl.load_workbook(output_path)
        try:
            ws = wb["成绩"]
            headers = [cell.value for cell in ws[1]]
            self.assertIn("一、客观题", headers)
            self.assertIn("二、程序分析", headers)
            self.assertIn("三、编程题", headers)
            row_values = [cell.value for cell in ws[2]]
            header_to_value = dict(zip(headers, row_values))
            self.assertEqual(header_to_value["Q1"], 2)
            self.assertEqual(header_to_value["Q2"], 2)
            self.assertEqual(header_to_value["Q3"], 5)
            self.assertEqual(header_to_value["Q4"], 4)
            self.assertEqual(header_to_value["一、客观题"], 4)
            self.assertEqual(header_to_value["二、程序分析"], 5)
            self.assertEqual(header_to_value["三、编程题"], 4)
            self.assertEqual(header_to_value["总分"], 13)
        finally:
            wb.close()
            del ws
    def test_llm_request_hash_changes_when_api_base_changes(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(build_llm_requests, "build_llm_requests 未实现")
        config = load_exam_config(self.config_path)
        config.setdefault("llm", {})["provider"] = "openai"
        config["llm"]["model"] = "demo-model"
        config["llm"]["api_base"] = "https://api.example.com/v1"
        first_hash = build_llm_requests(config)[0]["request_hash"]
        config["llm"]["api_base"] = "https://mirror.example.com/v1"
        second_hash = build_llm_requests(config)[0]["request_hash"]
        self.assertNotEqual(first_hash, second_hash)

    def test_normalize_llm_config_supports_legacy_openai_mode_and_defaults(self):
        self.assertIsNotNone(llm_grade_module, "llm_grade 模块未实现")
        config = {"llm": {"mode": "openai"}}
        normalized = llm_grade_module.normalize_llm_config(config)
        self.assertEqual(normalized["llm"]["mode"], "llm_api")
        self.assertEqual(normalized["llm"]["provider"], "openai")
        self.assertEqual(normalized["llm"]["api_key_env"], "OPENAI_API_KEY")
        self.assertEqual(normalized["llm"]["api_base"], "")
        self.assertEqual(normalized["llm"]["headers"], {})

    def test_call_llm_api_requires_named_api_key_env(self):
        self.assertIsNotNone(llm_grade_module, "llm_grade 模块未实现")
        llm_cfg = {
            "provider": "openai",
            "model": "demo-model",
            "api_key_env": "CUSTOM_EXAM_API_KEY"
        }
        with mock.patch.dict(os.environ, {}, clear=True):
            with self.assertRaisesRegex(RuntimeError, "CUSTOM_EXAM_API_KEY"):
                llm_grade_module.call_llm_api("prompt", llm_cfg)

        self.assertIsNotNone(grade_exam_module, "grade_exam 模块未实现")
        config = load_exam_config(self.config_path)
        requests_path = os.path.join(self.tempdir, "llm_requests.jsonl")
        config["files"]["llm_requests_jsonl"] = requests_path
        with open(requests_path, "w", encoding="utf-8") as f:
            f.write("{}\n")
        with open(self.llm_cache_path, "w", encoding="utf-8") as f:
            f.write("{}\n")
        removed, failures = grade_exam_module.cleanup_intermediate_files(config)
        self.assertCountEqual(removed, [requests_path, self.llm_path, self.llm_cache_path])
        self.assertEqual(failures, [])
        self.assertFalse(os.path.exists(requests_path))
        self.assertFalse(os.path.exists(self.llm_path))
        self.assertFalse(os.path.exists(self.llm_cache_path))
        self.assertFalse(os.path.exists(self.output_path))

    def test_grade_exam_cleanup_intermediate_files_skips_final_output(self):
        self.assertIsNotNone(grade_exam_module, "grade_exam 模块未实现")
        config = load_exam_config(self.config_path)
        config["files"]["llm_grades_jsonl"] = self.output_path
        with open(self.output_path, "w", encoding="utf-8") as f:
            f.write("final")
        removed, failures = grade_exam_module.cleanup_intermediate_files(config)
        self.assertNotIn(self.output_path, removed)
        self.assertEqual(failures, [])
        self.assertTrue(os.path.exists(self.output_path))

    def test_agent_runner_mode_prepare_exports_requests_without_calling_openai(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(llm_grade_module, "llm_grade 模块未实现")
        config = load_exam_config(self.config_path)
        config.setdefault("llm", {})["mode"] = "agent_runner"
        config["llm"]["agent_backend"] = "claude"
        with open(self.config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        requests_output = os.path.join(self.tempdir, "llm_requests.jsonl")
        with mock.patch("sys.argv", [
            "llm_grade.py",
            "--config", self.config_path,
            "--mode", "agent_runner",
            "--requests-output", requests_output,
        ]), mock.patch.dict(os.environ, {"BATCH_EXAM_GRADING_TEST": "1"}):
            llm_grade_module.main()
        self.assertTrue(os.path.exists(requests_output))
        with open(requests_output, encoding="utf-8") as f:
            rows = [json.loads(line) for line in f if line.strip()]
        self.assertEqual([item["question_id"] for item in rows], [3, 4])
        self.assertTrue(all("request_hash" in item for item in rows))

    def test_legacy_claude_agent_mode_is_normalized_to_agent_runner(self):
        self.assertIsNotNone(load_exam_config, "load_exam_config 未实现")
        self.assertIsNotNone(llm_grade_module, "llm_grade 模块未实现")
        config = load_exam_config(self.config_path)
        config.setdefault("llm", {})["mode"] = "claude_agent"
        with open(self.config_path, "w", encoding="utf-8") as f:
            json.dump(config, f, ensure_ascii=False, indent=2)
        requests_output = os.path.join(self.tempdir, "legacy_llm_requests.jsonl")
        with mock.patch("sys.argv", [
            "llm_grade.py",
            "--config", self.config_path,
            "--mode", "claude_agent",
            "--requests-output", requests_output,
        ]), mock.patch.dict(os.environ, {"BATCH_EXAM_GRADING_TEST": "1"}):
            llm_grade_module.main()
        self.assertTrue(os.path.exists(requests_output))


if __name__ == "__main__":
    unittest.main()
