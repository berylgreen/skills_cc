# -*- coding: utf-8 -*-
import csv
import hashlib
import json
import os
import re
from typing import Any, Dict, Iterable, List, Optional, Tuple

from docx import Document
import openpyxl
from openpyxl.styles import Alignment, Font, Border, Side
from openpyxl.utils import get_column_letter


DEFAULT_LLM_REVIEW_SHEET = "模型评分审计"


def _normalize_files_config(config: Dict[str, Any]) -> Dict[str, Any]:
    files_cfg = config.setdefault("files", {})
    if files_cfg.get("roster_file"):
        return config
    if files_cfg.get("roster_csv"):
        files_cfg["roster_file"] = files_cfg["roster_csv"]
    return config


def load_exam_config(path: str) -> Dict[str, Any]:
    with open(path, encoding="utf-8-sig") as f:
        config = json.load(f)
    config = _normalize_files_config(config)
    config["questions"] = sorted(config.get("questions", []), key=lambda item: int(item["id"]))
    return config


def question_map(config: Dict[str, Any]) -> Dict[int, Dict[str, Any]]:
    return {int(item["id"]): item for item in config.get("questions", [])}


def subjective_questions(config: Dict[str, Any]) -> List[Dict[str, Any]]:
    return [q for q in config.get("questions", []) if q.get("llm", {}).get("enabled")]


def llm_required(config: Dict[str, Any]) -> bool:
    return config.get("llm", {}).get("require_for_subjective", True)


def _strip_prefix(text: str) -> str:
    return re.sub(r"^\d+\s*[、\.．]\s*", "", text).strip()


def _find_section_lines(paragraphs: List[str], start_keywords: List[str], end_keywords: List[str]) -> List[str]:
    active = False
    lines = []
    for text in paragraphs:
        if any(keyword in text for keyword in start_keywords):
            active = True
            continue
        if active and any(keyword in text for keyword in end_keywords):
            break
        if active:
            lines.append(text)
    return lines


def _extract_choice_answers(paragraphs: List[str], config: Dict[str, Any]) -> Dict[int, str]:
    section_cfg = config.get("parsing", {}).get("sections", {}).get("choice", {})
    lines = _find_section_lines(
        paragraphs,
        section_cfg.get("start_keywords", ["单选题", "选择题"]),
        section_cfg.get("end_keywords", ["填空题", "程序分析"]),
    )
    answers = {}
    for line in lines:
        for qn, ans in re.findall(r"(\d+)\s*[、\.．]\s*([A-Za-z]+)", line):
            answers[int(qn)] = ans.upper()
    return answers


def _extract_answer_by_pattern(paragraphs: List[str], pattern: str) -> str:
    regex = re.compile(pattern)
    for line in paragraphs:
        match = regex.search(line)
        if match:
            return match.group(1).strip() if match.groups() else match.group(0).strip()
    return ""


def _extract_answer_by_pattern_in_section(
    paragraphs: List[str],
    config: Dict[str, Any],
    section_name: str,
    pattern: str,
) -> str:
    section_cfg = config.get("parsing", {}).get("sections", {}).get(section_name, {})
    lines = _find_section_lines(
        paragraphs,
        section_cfg.get("start_keywords", []),
        section_cfg.get("end_keywords", []),
    )
    return _extract_answer_by_pattern(lines, pattern)


def _extract_section_question(paragraphs: List[str], config: Dict[str, Any], question: Dict[str, Any]) -> str:
    extract_cfg = question.get("extract", {})
    section_cfg = config.get("parsing", {}).get("sections", {}).get("analysis", {})
    lines = _find_section_lines(
        paragraphs,
        section_cfg.get("start_keywords", ["程序分析", "分析题"]),
        section_cfg.get("end_keywords", ["编程题", "程序设计"]),
    )
    q_number = int(extract_cfg.get("question_number", 1))
    pattern = re.compile(
        rf"^[{q_number}{_to_fullwidth_digit(q_number)}]\s*[\.．]\s*[（(].*?[）)]\s*(.*)"
    )
    parts: List[str] = []
    active = False
    for line in lines:
        match = pattern.match(line)
        if match:
            active = True
            rest = match.group(1).strip()
            if rest:
                parts.append(rest)
            continue
        if active and re.match(r"^[0-9０-９]\s*[\.．]\s*[（(].*?[）)]", line):
            break
        if active and line:
            parts.append(line)
    return " ".join(parts).strip()


def _to_fullwidth_digit(number: int) -> str:
    return str(number).translate(str.maketrans("0123456789", "０１２３４５６７８９"))


def _extract_code_tables(doc: Document, config: Dict[str, Any]) -> List[str]:
    parsing = config.get("parsing", {})
    keywords = parsing.get("code_table_keywords", ["class", "public", "static", "void"])
    min_length = int(parsing.get("code_min_length", 20))
    tables = []
    for table in doc.tables:
        text = "\n".join(cell.text.strip() for row in table.rows for cell in row.cells if cell.text.strip())
        if len(text) >= min_length and any(keyword in text for keyword in keywords):
            tables.append(text)
    return tables


def extract_student_answers(config: Dict[str, Any], docx_path: str) -> Dict[int, str]:
    doc = Document(docx_path)
    paragraphs = [p.text.strip() for p in doc.paragraphs]
    choice_answers = _extract_choice_answers(paragraphs, config)
    code_tables = _extract_code_tables(doc, config)

    answers: Dict[int, str] = {}
    for question in config.get("questions", []):
        qid = int(question["id"])
        qtype = question.get("type", "")
        extract_cfg = question.get("extract", {})
        source = extract_cfg.get("source", "")
        if source == "choice_inline" and qtype == "choice":
            answers[qid] = choice_answers.get(qid, "")
        elif source == "choice_inline" and qtype == "fill":
            pattern = extract_cfg.get("answer_pattern")
            if pattern:
                answers[qid] = _extract_answer_by_pattern(paragraphs, pattern)
            else:
                answers[qid] = choice_answers.get(qid, "")
        elif source == "section_question":
            answers[qid] = _extract_section_question(paragraphs, config, question)
        elif source == "code_table":
            table_index = int(extract_cfg.get("table_index", 0))
            answers[qid] = code_tables[table_index] if table_index < len(code_tables) else ""
        elif source == "regex":
            answers[qid] = _extract_answer_by_pattern(paragraphs, extract_cfg.get("pattern", ""))
        elif source == "section_regex":
            answers[qid] = _extract_answer_by_pattern_in_section(
                paragraphs,
                config,
                extract_cfg.get("section_name", ""),
                extract_cfg.get("pattern", ""),
            )
        else:
            answers[qid] = ""
    return answers


def extract_sid_from_filename(config: Dict[str, Any], filename: str) -> Tuple[Optional[str], Optional[str]]:
    pattern = config.get("parsing", {}).get("student_id_pattern", r"(\d{10,12})")
    match = re.search(pattern, filename)
    if not match:
        try:
            decoded_pattern = pattern.encode("utf-8").decode("unicode_escape")
            match = re.search(decoded_pattern, filename)
        except UnicodeDecodeError:
            match = None
    if not match:
        return None, None
    sid = match.group(1)
    rest = filename[match.end():]
    rest = re.sub(r"\.docx$", "", rest, flags=re.I)
    rest = re.sub(r"_?答题纸.*$", "", rest)
    rest = re.sub(r"^[_\-\s]+", "", rest)
    rest = re.sub(r"[_\-\s]+$", "", rest)
    return sid, rest.strip()


def _score_exact(answer: str, expected: str, score: float) -> float:
    return score if answer.strip().upper() == str(expected).strip().upper() else 0.0


def _score_contains_any(answer: str, answers: List[str], score: float) -> float:
    lowered = answer.lower()
    return score if any(item.lower() in lowered for item in answers) else 0.0


def _score_regex_any(answer: str, patterns: List[str], score: float) -> float:
    return score if any(re.search(pattern, answer, re.I) for pattern in patterns) else 0.0


def _score_numeric_map(answer: str, mapping: Dict[str, Any]) -> float:
    cleaned = re.sub(r"^(答|解答|答案)\s*[：:]\s*", "", answer.strip())
    match = re.search(r"\d+", cleaned)
    if match:
        return float(mapping.get(match.group(0), 0))
    return 0.0


def _score_keyword_points(answer: str, points: List[Dict[str, Any]]) -> float:
    score = 0.0
    lowered = answer.lower()
    for point in points:
        keywords = point.get("keywords", [])
        if all(keyword.lower() in lowered for keyword in keywords):
            score += float(point.get("score", 0))
    return score


def grade_answer(question: Dict[str, Any], answer: str, use_fallback: bool = False) -> float:
    grading = question.get("fallback_grading" if use_fallback else "grading", {})
    if not grading:
        return 0.0
    mode = grading.get("mode", "")
    max_score = float(question.get("score", grading.get("score", 0)))
    if answer is None:
        answer = ""
    answer = str(answer)
    if mode == "exact":
        score = _score_exact(answer, grading.get("answer", ""), max_score)
    elif mode == "contains_any":
        score = _score_contains_any(answer, grading.get("answers", []), max_score)
    elif mode == "regex_any":
        score = _score_regex_any(answer, grading.get("patterns", []), max_score)
    elif mode == "numeric_map":
        score = _score_numeric_map(answer, grading.get("mapping", {}))
    elif mode == "keyword_points":
        score = _score_keyword_points(answer, grading.get("points", []))
    else:
        score = 0.0
    return min(max(score, 0.0), max_score)


def load_llm_grades(path: str) -> Tuple[Dict[Tuple[str, int], Dict[str, Any]], List[Dict[str, Any]]]:
    mapping: Dict[Tuple[str, int], Dict[str, Any]] = {}
    rows: List[Dict[str, Any]] = []
    if not path or not os.path.exists(path):
        return mapping, rows
    with open(path, encoding="utf-8-sig") as f:
        for line in f:
            line = line.strip()
            if not line:
                continue
            item = json.loads(line)
            sid = str(item.get("student_id", "")).strip()
            qid = int(item.get("question_id"))
            mapping[(sid, qid)] = item
            rows.append(item)
    return mapping, rows


def _normalize_score(value: Any) -> Any:
    score = float(value)
    return int(score) if score.is_integer() else score


def build_request_hash(config: Dict[str, Any], request: Dict[str, Any]) -> str:
    payload = {
        "provider": config.get("llm", {}).get("provider", "openai"),
        "model": config.get("llm", {}).get("model", ""),
        "question_id": request["question_id"],
        "question_type": request["question_type"],
        "max_score": request["max_score"],
        "prompt": request["prompt"],
        "reference_answer": request["reference_answer"],
        "rubric": request["rubric"],
        "student_answer": request["student_answer"],
    }
    text = json.dumps(payload, ensure_ascii=False, sort_keys=True)
    return hashlib.sha256(text.encode("utf-8")).hexdigest()


def build_llm_requests(config: Dict[str, Any]) -> List[Dict[str, Any]]:
    requests: List[Dict[str, Any]] = []
    answer_folder = config.get("files", {}).get("answer_folder", "")
    for filename in sorted(os.listdir(answer_folder)):
        if not filename.lower().endswith(".docx"):
            continue
        sid, _ = extract_sid_from_filename(config, filename)
        if not sid:
            continue
        answers = extract_student_answers(config, os.path.join(answer_folder, filename))
        for question in subjective_questions(config):
            qid = int(question["id"])
            requests.append({
                "student_id": sid,
                "filename": filename,
                "question_id": qid,
                "question_type": question.get("type", ""),
                "max_score": question.get("score", 0),
                "prompt": question.get("llm", {}).get("prompt", ""),
                "reference_answer": question.get("llm", {}).get("reference_answer", ""),
                "rubric": question.get("llm", {}).get("rubric", []),
                "student_answer": answers.get(qid, "")
            })
            requests[-1]["request_hash"] = build_request_hash(config, requests[-1])
    return requests


def read_roster(config: Dict[str, Any]) -> List[Dict[str, str]]:
    roster_cfg = config.get("roster", {})
    path = config.get("files", {}).get("roster_file", "")
    if path.lower().endswith(".xlsx"):
        return _read_roster_xlsx(path, roster_cfg)
    rows: List[Dict[str, str]] = []
    with open(path, encoding="utf-8-sig", newline="") as f:
        reader = csv.DictReader(f)
        for row in reader:
            rows.append({
                "student_id": str(row.get(roster_cfg.get("student_id_field", "student_id"), "")).strip(),
                "name": str(row.get(roster_cfg.get("name_field", "name"), "")).strip(),
                "seq": str(row.get(roster_cfg.get("seq_field", "seq"), "")).strip()
            })
    return rows


def _read_roster_xlsx(path: str, roster_cfg: Dict[str, str]) -> List[Dict[str, str]]:
    wb = openpyxl.load_workbook(path, read_only=True)
    ws = wb.active
    rows_iter = ws.iter_rows(values_only=True)
    headers = [str(value).strip() if value is not None else "" for value in next(rows_iter)]
    rows: List[Dict[str, str]] = []
    student_id_field = roster_cfg.get("student_id_field", "student_id")
    name_field = roster_cfg.get("name_field", "name")
    seq_field = roster_cfg.get("seq_field", "seq")
    for values in rows_iter:
        raw = {headers[idx]: values[idx] for idx in range(min(len(headers), len(values)))}
        rows.append({
            "student_id": str(raw.get(student_id_field, "")).strip(),
            "name": str(raw.get(name_field, "")).strip(),
            "seq": str(raw.get(seq_field, "")).strip()
        })
    wb.close()
    return rows


def _section_groups(config: Dict[str, Any]) -> Dict[str, List[int]]:
    groups: Dict[str, List[int]] = {}
    for question in config.get("questions", []):
        section = question.get("section", "未分组")
        groups.setdefault(section, []).append(int(question["id"]))
    return groups


def grade_roster(config: Dict[str, Any]) -> Dict[str, Any]:
    llm_path = config.get("files", {}).get("llm_grades_jsonl", "")
    llm_mapping, llm_records = load_llm_grades(llm_path)
    llm_requests = build_llm_requests(config)
    llm_request_map = {(item["student_id"], int(item["question_id"])): item for item in llm_requests}
    answer_folder = config.get("files", {}).get("answer_folder", "")
    question_by_id = question_map(config)
    extracted: Dict[str, Dict[int, str]] = {}
    for filename in sorted(os.listdir(answer_folder)):
        if not filename.lower().endswith(".docx"):
            continue
        sid, _ = extract_sid_from_filename(config, filename)
        if not sid:
            continue
        extracted[sid] = extract_student_answers(config, os.path.join(answer_folder, filename))

    roster = read_roster(config)
    scores: Dict[str, Dict[str, Any]] = {}
    for student in roster:
        sid = student["student_id"]
        answers = extracted.get(sid, {})
        student_scores: Dict[str, Any] = {}
        for qid, question in question_by_id.items():
            llm_enabled = question.get("llm", {}).get("enabled", False)
            if llm_enabled and (sid, qid) in llm_mapping:
                row = llm_mapping[(sid, qid)]
                request = llm_request_map.get((sid, qid))
                if request and row.get("request_hash") and row.get("request_hash") != request["request_hash"]:
                    raise ValueError(f"主观题模型评分已过期: {sid} Q{qid}")
                if request and not row.get("request_hash"):
                    raise ValueError(f"主观题模型评分缺少 request_hash: {sid} Q{qid}")
                score = _normalize_score(row["score"])
            elif llm_enabled:
                if llm_required(config):
                    raise ValueError(f"缺少主观题模型评分: {sid} Q{qid}")
                score = _normalize_score(grade_answer(question, answers.get(qid, ""), use_fallback=True))
            else:
                score = _normalize_score(grade_answer(question, answers.get(qid, "")))
            student_scores[f"q{qid}"] = score
        scores[sid] = student_scores

    return {
        "roster": roster,
        "scores": scores,
        "llm_records": llm_records,
        "section_groups": _section_groups(config)
    }


def write_excel(config: Dict[str, Any], result: Dict[str, Any]) -> str:
    output_path = config.get("files", {}).get("output_xlsx", "成绩表.xlsx")
    questions = config.get("questions", [])
    question_ids = [int(item["id"]) for item in questions]
    section_groups = result["section_groups"]

    wb = openpyxl.Workbook()
    ws = wb.active
    ws.title = "成绩"

    headers = ["序号", "学号", "姓名"] + [f"Q{qid}" for qid in question_ids] + ["总分"] + list(section_groups.keys())
    thin = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin"))

    for col, header in enumerate(headers, 1):
        cell = ws.cell(1, col, header)
        cell.font = Font(bold=True, size=10)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin

    for row_idx, student in enumerate(result["roster"], 2):
        sid = student["student_id"]
        ws.cell(row_idx, 1, int(student["seq"]) if student["seq"].isdigit() else row_idx - 1)
        ws.cell(row_idx, 2, sid)
        ws.cell(row_idx, 3, student["name"])
        total = 0.0
        for offset, qid in enumerate(question_ids, 4):
            score = result["scores"].get(sid, {}).get(f"q{qid}", 0)
            ws.cell(row_idx, offset, score)
            total += float(score)
        total_value = int(total) if float(total).is_integer() else total
        ws.cell(row_idx, 4 + len(question_ids), total_value)
        for section_idx, (section_name, ids) in enumerate(section_groups.items(), 5 + len(question_ids)):
            section_total = sum(float(result["scores"].get(sid, {}).get(f"q{qid}", 0)) for qid in ids)
            ws.cell(row_idx, section_idx, int(section_total) if section_total.is_integer() else section_total)
        for col in range(1, len(headers) + 1):
            ws.cell(row_idx, col).border = thin
            ws.cell(row_idx, col).alignment = Alignment(horizontal="center", vertical="center")

    ws.freeze_panes = "D2"
    _write_llm_audit_sheet(wb, result["llm_records"], {item["student_id"]: item["name"] for item in result["roster"]})
    wb.save(output_path)
    return output_path


def _write_llm_audit_sheet(wb, llm_records: List[Dict[str, Any]], roster_names: Dict[str, str]) -> None:
    if not llm_records:
        return
    ws = wb.create_sheet(DEFAULT_LLM_REVIEW_SHEET)
    headers = ["学号", "姓名", "题号", "题型", "得分", "满分", "置信度", "需人工复核", "扣分点", "依据", "文件名"]
    thin = Border(left=Side(style="thin"), right=Side(style="thin"), top=Side(style="thin"), bottom=Side(style="thin"))
    for col, header in enumerate(headers, 1):
        cell = ws.cell(1, col, header)
        cell.font = Font(bold=True, size=10)
        cell.alignment = Alignment(horizontal="center", vertical="center")
        cell.border = thin
    for row_idx, item in enumerate(llm_records, 2):
        row = [
            item.get("student_id", ""),
            roster_names.get(str(item.get("student_id", "")), ""),
            item.get("question_id", ""),
            item.get("question_type", ""),
            item.get("score", ""),
            item.get("max_score", ""),
            item.get("confidence", ""),
            "是" if item.get("needs_human_review") else "否",
            json.dumps(item.get("deductions", []), ensure_ascii=False),
            json.dumps(item.get("evidence", []), ensure_ascii=False),
            item.get("filename", "")
        ]
        for col, value in enumerate(row, 1):
            cell = ws.cell(row_idx, col, value)
            cell.border = thin
            cell.alignment = Alignment(horizontal="center" if col <= 8 else "left", vertical="top", wrap_text=True)
    ws.freeze_panes = "A2"
    widths = [16, 10, 6, 10, 6, 6, 10, 12, 40, 40, 28]
    for idx, width in enumerate(widths, 1):
        ws.column_dimensions[get_column_letter(idx)].width = width
