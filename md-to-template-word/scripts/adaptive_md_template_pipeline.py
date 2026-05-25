#!/usr/bin/env python
"""Starter pipeline for adaptive DOCX template reconstruction.

This script is intentionally not a universal converter. It gives an AI agent a
safe base to copy and modify for a user's concrete template and source project.
For institutional templates, running this file unchanged is only a smoke test:
it appends the rough body to the template and does not delete sample pages, fill
cover fields, replace abstracts, rebuild TOC entries, or repair sections.

Typical use:
  1. Convert LaTeX/Markdown/PDF into a rough body.docx.
  2. Inspect template.docx with inspect_docx_template.py.
  3. Copy this script into the work directory and patch mapping rules.
  4. Run it to produce final.docx, then finalize with Word COM.

Configuration:
  Pass --config config.json to override defaults. Recognized keys:

  body_style                 Target style name for body paragraphs (e.g. "论文正文").
  body_candidate_styles      Source style names treated as "body candidates" and
                             remapped to body_style. Defaults cover both English
                             ("Normal", "Body Text") and Chinese ("正文") names.
  unnumbered_h1              List of heading texts whose Heading 1 numbering
                             should be suppressed (e.g. 摘要 / 参考文献).
  unnumbered_heading_styles  Style names treated as "Heading 1" — defaults cover
                             both "Heading 1" and "标题 1".
  caption_regex              Regex matched against paragraph.text for captions.
                             Default matches both 图/表/Figure/Table numbering.
  body_font_name             Latin font name for body runs. None = leave as-is.
  body_east_asia_font        East-Asian font name for body runs. None = leave as-is.
  body_font_size_pt          Body run font size in points. None = leave as-is.
  table_font_name            Same, for table cells (only used when three-line
                             tables are enabled).
  table_east_asia_font       Same, for table cells.
  table_font_size_pt         Same, for table cells.
  enable_three_line_tables   Coerce all tables into three-line style. Off by
                             default — turn on only when your template really
                             requires it. CLI: --three-line-tables.
  enable_black_hyperlinks    Force hyperlinks to black / no underline. On by
                             default for print-style thesis output. CLI:
                             --keep-hyperlink-color disables it.
  remap_styles_by_name      Remap copied source style ids to template style ids
                             with the same visible style name. On by default;
                             this prevents Heading 1/2/3 from turning into an
                             unrelated template style when style ids collide.
  formatting_start_marker   Only apply global body/caption/table/hyperlink
                             formatting at or after this normalized marker.
                             Use this to protect native cover/declaration pages.
  formatting_end_marker     Optional normalized marker where global formatting
                             stops.
  formatting_include_start_marker
                             Whether the start marker paragraph is in scope.
                             Default true.
  clear_header_references   Remove section header references after reconstruction.
                             Use when deleting template sample sections leaves
                             a back-matter header such as 致谢 on body pages.
  clear_footer_references   Remove section footer references after reconstruction.
  heading_1_page_breaks     Force Heading 1 paragraphs to start on a new page.

  See presets/ for ready-to-use config samples (e.g. zhengzhou_thesis.json).
"""
from __future__ import annotations

import argparse
import copy
import json
import re
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


def clear_document_body(doc: Document) -> None:
    body = doc.element.body
    sect_pr = body.sectPr
    for child in list(body):
        if child.tag != qn("w:sectPr"):
            body.remove(child)
    if sect_pr is not None and body.sectPr is None:
        body.append(sect_pr)


def parse_markdown_table(table_lines: list[str]) -> list[list[str]]:
    rows: list[list[str]] = []
    for idx, line in enumerate(table_lines):
        if idx == 1:
            parts = [part.strip() for part in line.strip().strip("|").split("|")]
            if parts and all(re.fullmatch(r"[:\-\s]+", part or "-") for part in parts):
                continue
        rows.append([cell.strip() for cell in line.strip().strip("|").split("|")])
    return rows


def parse_markdown_blocks(md_text: str) -> list[dict]:
    lines = md_text.splitlines()
    blocks: list[dict] = []
    i = 0
    title_seen = False
    while i < len(lines):
        stripped = lines[i].strip()
        if not stripped:
            i += 1
            continue

        heading_match = re.match(r"^(#{1,6})\s+(.*)$", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            if level == 1 and not title_seen:
                blocks.append({"type": "title", "text": text})
                title_seen = True
            else:
                blocks.append({"type": "heading", "level": level, "text": text})
            i += 1
            continue

        if stripped.startswith("|"):
            table_lines: list[str] = []
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            blocks.append({"type": "table", "rows": parse_markdown_table(table_lines)})
            continue

        if re.match(r"^\s*[-*]\s+", lines[i]):
            items: list[dict] = []
            while i < len(lines) and re.match(r"^\s*[-*]\s+", lines[i]):
                line = lines[i]
                indent = len(line) - len(line.lstrip())
                level = indent // 2
                if level > 8:
                    level = 8
                text = re.sub(r"^\s*[-*]\s+", "", line)
                items.append({"level": level, "text": text.strip()})
                i += 1
            blocks.append({"type": "list", "items": items})
            continue

        if re.match(r"^[-*_]{3,}$", stripped):
            blocks.append({"type": "hr"})
            i += 1
            continue

        para_lines = [stripped]
        i += 1
        while i < len(lines):
            next_line = lines[i].strip()
            if not next_line:
                i += 1
                break
            if (
                re.match(r"^(#{1,6})\s+", next_line)
                or next_line.startswith("|")
                or re.match(r"^[-*]\s+", next_line)
            ):
                break
            para_lines.append(next_line)
            i += 1
        blocks.append({"type": "paragraph", "text": "".join(para_lines)})

    return blocks


def strip_manual_heading_number(text: str, level: int) -> str:
    stripped = text.strip()
    patterns = {
        2: r'^[一二三四五六七八九十]+、\s*',
        3: r'^\d+\.\s*',
        4: r'^[（(]\d+[)）]\s*',
    }
    pattern = patterns.get(level)
    if not pattern:
        return stripped
    return re.sub(pattern, '', stripped).strip() or stripped


# Default style names cover both English and Chinese localized templates.
DEFAULT_UNNUMBERED_HEADING_STYLES = ("Heading 1", "标题 1")
DEFAULT_BODY_CANDIDATE_STYLES = ("Normal", "Body Text", "正文", "标准")
# Generic caption pattern: 图/表/Figure/Fig./Table/Tab. + number, with optional
# section-style numbering like "3.1" / "3-1".
DEFAULT_CAPTION_REGEX = r"^(图|表|Figure|Fig\.|Table|Tab\.)\s*\d+([.\-]\d+)?\s+"


def normalize_marker_text(text: str | None) -> str:
    return re.sub(r"\s+", "", text or "")


def body_element_text(element) -> str:
    return "".join(t.text or "" for t in element.iter(qn("w:t")))


def body_elements_in_marker_scope(
    doc: Document,
    *,
    start_marker: str | None = None,
    end_marker: str | None = None,
    include_start_marker: bool = True,
):
    """Yield top-level body elements inside a marker-delimited formatting scope.

    Institutional templates often have cover/declaration pages before the real
    abstract/body. Global style passes must be scoped so those native template
    pages remain byte-for-byte close to the template except intentional text
    replacements.
    """
    normalized_start = normalize_marker_text(start_marker)
    normalized_end = normalize_marker_text(end_marker)
    in_scope = not normalized_start
    found_start = not normalized_start

    for child in doc.element.body:
        child_text = normalize_marker_text(body_element_text(child))
        if not in_scope:
            if normalized_start and normalized_start in child_text:
                in_scope = True
                found_start = True
                if include_start_marker:
                    yield child
                continue
        else:
            if normalized_end and normalized_end in child_text:
                break
            yield child

    if normalized_start and not found_start:
        raise ValueError(f"formatting_start_marker not found: {start_marker}")


def iter_paragraphs_in_marker_scope(
    doc: Document,
    *,
    start_marker: str | None = None,
    end_marker: str | None = None,
    include_start_marker: bool = True,
):
    allowed = {
        id(element)
        for element in body_elements_in_marker_scope(
            doc,
            start_marker=start_marker,
            end_marker=end_marker,
            include_start_marker=include_start_marker,
        )
        if element.tag == qn("w:p")
    }
    for paragraph in doc.paragraphs:
        if id(paragraph._p) in allowed:
            yield paragraph


def table_ids_in_marker_scope(
    doc: Document,
    *,
    start_marker: str | None = None,
    end_marker: str | None = None,
    include_start_marker: bool = True,
) -> set[int] | None:
    if not start_marker and not end_marker:
        return None
    return {
        id(element)
        for element in body_elements_in_marker_scope(
            doc,
            start_marker=start_marker,
            end_marker=end_marker,
            include_start_marker=include_start_marker,
        )
        if element.tag == qn("w:tbl")
    }


def replace_paragraph_text_preserving_runs(paragraph, text: str) -> None:
    """Replace visible paragraph text without deleting paragraph/run formatting."""
    text_runs = [run for run in paragraph.runs if run.text]
    if not text_runs:
        paragraph.add_run(text)
        return
    first = text_runs[0]
    first.text = text
    for run in paragraph.runs:
        if run is not first and run.text:
            run.text = ""


def replace_text_preserving_format(
    paragraph,
    needle: str,
    replacement: str,
    *,
    compact_match: bool = False,
) -> bool:
    """Replace a placeholder while preserving the template paragraph shape.

    Prefer exact in-run replacement. If the placeholder spans multiple runs or
    uses spacing that needs compact matching, keep the existing runs and only
    move the replacement text into the first text run.
    """
    if not needle:
        return False
    for run in paragraph.runs:
        if needle in run.text:
            run.text = run.text.replace(needle, replacement)
            return True
    if compact_match and normalize_marker_text(needle) in normalize_marker_text(paragraph.text):
        replace_paragraph_text_preserving_runs(paragraph, replacement)
        return True
    return False


def set_east_asia_font(run, font_name: str) -> None:
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), font_name)
def clear_paragraph_indent(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    ind = ppr.find(qn("w:ind"))
    if ind is None:
        ind = OxmlElement("w:ind")
        ppr.append(ind)
    for attr in list(ind.attrib):
        del ind.attrib[attr]
    ind.set(qn("w:left"), "0")
    ind.set(qn("w:firstLine"), "0")
    ind.set(qn("w:leftChars"), "0")
    ind.set(qn("w:firstLineChars"), "0")



def write_text_paragraph(
    doc: Document,
    text: str,
    *,
    style_name: str,
    align: int | None = None,
    bold: bool | None = None,
    font_size_pt: float | None = None,
    east_asia_font: str | None = None,
) -> object:
    paragraph = doc.add_paragraph(style=style_name)
    if align is not None:
        paragraph.alignment = align
        paragraph.paragraph_format.first_line_indent = None
    parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__)', text or "")
    for part in parts:
        if not part:
            continue
        is_bold_part = False
        run_text = part
        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
            is_bold_part = True
            run_text = part[2:-2]
        elif part.startswith('__') and part.endswith('__') and len(part) >= 4:
            is_bold_part = True
            run_text = part[2:-2]
            
        run_text = run_text.replace("`", "")
        run = paragraph.add_run(run_text)
        if bold is not None:
            run.bold = bold
        elif is_bold_part:
            run.bold = True
            
        if font_size_pt is not None:
            run.font.size = Pt(font_size_pt)
        if east_asia_font is not None:
            set_east_asia_font(run, east_asia_font)
    return paragraph


def set_paragraph_numbering(paragraph, *, num_id: int, ilvl: int) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl_node = num_pr.find(qn("w:ilvl"))
    if ilvl_node is None:
        ilvl_node = OxmlElement("w:ilvl")
        num_pr.append(ilvl_node)
    ilvl_node.set(qn("w:val"), str(ilvl))
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def set_list_level(paragraph, ilvl: int, num_id: int = 2) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl_node = num_pr.find(qn("w:ilvl"))
    if ilvl_node is None:
        ilvl_node = OxmlElement("w:ilvl")
        num_pr.append(ilvl_node)
    ilvl_node.set(qn("w:val"), str(ilvl))
    num_id_node = num_pr.find(qn("w:numId"))
    if num_id_node is None:
        num_id_node = OxmlElement("w:numId")
        num_pr.append(num_id_node)
    num_id_node.set(qn("w:val"), str(num_id))


def add_page_break_paragraph(doc: Document) -> None:
    paragraph = doc.add_paragraph()
    paragraph.add_run().add_break(WD_BREAK.PAGE)


def estimate_text_units(text: str) -> int:
    units = 0
    for ch in (text or "").strip():
        if ch.isspace():
            continue
        units += 2 if unicodedata.east_asian_width(ch) in {"F", "W", "A"} else 1
    return units


def compute_table_widths_dxa(rows: list[list[str]], col_count: int, *, total_dxa: int = 9070) -> list[int]:
    max_units = [0] * col_count
    for row in rows:
        for col_idx in range(col_count):
            cell_text = row[col_idx] if col_idx < len(row) else ""
            max_units[col_idx] = max(max_units[col_idx], estimate_text_units(cell_text))

    min_units = 8 if col_count <= 4 else 6
    max_cap = 36
    scores = [min(max(units, min_units), max_cap) for units in max_units]
    total_score = sum(scores) or col_count

    widths: list[int] = []
    used = 0
    for idx, score in enumerate(scores):
        if idx == col_count - 1:
            width = total_dxa - used
        else:
            width = max(int(round(total_dxa * score / total_score)), 900)
            used += width
        widths.append(width)
    return widths


def set_cell_width_dxa(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def apply_table_widths(table, widths_dxa: list[int]) -> None:
    table.autofit = False
    tbl = table._tbl
    tbl_pr = table._tbl.tblPr
    tbl_layout = tbl_pr.first_child_found_in("w:tblLayout")
    if tbl_layout is None:
        tbl_layout = OxmlElement("w:tblLayout")
        tbl_pr.append(tbl_layout)
    tbl_layout.set(qn("w:type"), "fixed")

    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))
    tbl_w.set(qn("w:type"), "dxa")

    tbl_grid = tbl.first_child_found_in("w:tblGrid")
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.insert(1, tbl_grid)
    for grid_col in list(tbl_grid):
        tbl_grid.remove(grid_col)
    for width_dxa in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width_dxa))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            set_cell_width_dxa(cell, widths_dxa[col_idx])


def rebuild_from_markdown(
    doc: Document,
    md_text: str,
    *,
    table_captions: list[str] | None = None,
    body_east_asia_font: str | None = None,
    insert_toc: bool = True,
) -> None:
    clear_document_body(doc)
    blocks = parse_markdown_blocks(md_text)
    table_captions = table_captions or []
    table_index = 0
    saw_body_heading = False
    toc_inserted = False

    for block in blocks:
        block_type = block["type"]
        if block_type == "title":
            write_text_paragraph(
                doc,
                block["text"],
                style_name="Normal",
                align=WD_ALIGN_PARAGRAPH.CENTER,
                bold=True,
                font_size_pt=18,
                east_asia_font=body_east_asia_font,
            )
            continue
        if block_type == "heading":
            if insert_toc and not toc_inserted:
                add_page_break_paragraph(doc)
                write_text_paragraph(
                    doc,
                    "目录",
                    style_name="Normal",
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    bold=True,
                    font_size_pt=16,
                    east_asia_font=body_east_asia_font,
                )
                add_page_break_paragraph(doc)
                toc_inserted = True
            saw_body_heading = True
            clean_text = strip_manual_heading_number(block["text"], block["level"])
            heading_style = {
                2: "Heading 1",
                3: "Heading 2",
                4: "Heading 3",
            }.get(block["level"], "Normal")
            paragraph = write_text_paragraph(
                doc,
                clean_text,
                style_name=heading_style,
                east_asia_font=body_east_asia_font,
            )
            heading_level = {2: 0, 3: 1, 4: 2}.get(block["level"])
            if heading_level is not None:
                set_paragraph_numbering(paragraph, num_id=1, ilvl=heading_level)
            continue
        if block_type == "hr":
            # Ignore markdown horizontal rules instead of rendering them as literal "---"
            continue
        if block_type == "paragraph":
            is_subtitle = not saw_body_heading and len(doc.paragraphs) == 1
            write_text_paragraph(
                doc,
                block["text"],
                style_name="Normal",
                align=WD_ALIGN_PARAGRAPH.CENTER if is_subtitle else None,
                east_asia_font=body_east_asia_font,
            )
            continue
        if block_type == "list":
            for item in block["items"]:
                paragraph = write_text_paragraph(
                    doc,
                    item["text"],
                    style_name="List Paragraph",
                    east_asia_font=body_east_asia_font,
                )
                set_list_level(paragraph, item["level"])
                if item["level"] > 0:
                    paragraph.paragraph_format.left_indent = Pt(36 + 21 * item["level"])
            continue
        if block_type == "table":
            if table_index < len(table_captions):
                write_text_paragraph(
                    doc,
                    table_captions[table_index],
                    style_name="Table Caption",
                    align=WD_ALIGN_PARAGRAPH.CENTER,
                    east_asia_font=body_east_asia_font,
                )
            rows = block["rows"]
            if not rows:
                continue
            col_count = max(len(row) for row in rows)
            table = doc.add_table(rows=len(rows), cols=col_count)
            widths_dxa = compute_table_widths_dxa(rows, col_count)
            apply_table_widths(table, widths_dxa)
            for row_idx, row in enumerate(rows):
                for col_idx in range(col_count):
                    cell_text = row[col_idx] if col_idx < len(row) else ""
                    cell = table.cell(row_idx, col_idx)
                    cell.text = ""
                    paragraph = cell.paragraphs[0]
                    paragraph.alignment = (
                        WD_ALIGN_PARAGRAPH.CENTER
                        if row_idx == 0 or len(cell_text) <= 18
                        else WD_ALIGN_PARAGRAPH.LEFT
                    )
                    clear_paragraph_indent(paragraph)
                    parts = re.split(r'(\*\*[^*]+\*\*|__[^_]+__)', cell_text or "")
                    for part in parts:
                        if not part:
                            continue
                        is_bold_part = False
                        run_text = part
                        if part.startswith('**') and part.endswith('**') and len(part) >= 4:
                            is_bold_part = True
                            run_text = part[2:-2]
                        elif part.startswith('__') and part.endswith('__') and len(part) >= 4:
                            is_bold_part = True
                            run_text = part[2:-2]
                            
                        run_text = run_text.replace("`", "")
                        run = paragraph.add_run(run_text)
                        if row_idx == 0 or is_bold_part:
                            run.bold = True
                        if body_east_asia_font is not None:
                            set_east_asia_font(run, body_east_asia_font)
            doc.add_paragraph("")
            table_index += 1


def suppress_heading_number(paragraph) -> None:
    ppr = paragraph._p.get_or_add_pPr()
    num_pr = ppr.find(qn("w:numPr"))
    if num_pr is None:
        num_pr = OxmlElement("w:numPr")
        ppr.append(num_pr)
    ilvl = num_pr.find(qn("w:ilvl"))
    if ilvl is None:
        ilvl = OxmlElement("w:ilvl")
        num_pr.append(ilvl)
    ilvl.set(qn("w:val"), "0")
    num_id = num_pr.find(qn("w:numId"))
    if num_id is None:
        num_id = OxmlElement("w:numId")
        num_pr.append(num_id)
    num_id.set(qn("w:val"), "0")


def build_style_id_map(target_doc: Document, source_doc: Document) -> dict[str, str]:
    """Map source style ids to target style ids by visible style name.

    python-docx exposes paragraph.style.name, but copied OOXML stores style ids
    such as Heading1. Institutional templates often reuse different ids for the
    same visible style names, so appending raw source XML can silently turn
    Heading 1 into a body style.
    """
    target_by_name = {style.name: style.style_id for style in target_doc.styles}
    mapping: dict[str, str] = {}
    for style in source_doc.styles:
        target_style_id = target_by_name.get(style.name)
        if target_style_id:
            mapping[style.style_id] = target_style_id
    return mapping


def remap_style_ids(element, style_id_map: dict[str, str]) -> None:
    if not style_id_map:
        return
    style_tags = {qn("w:pStyle"), qn("w:rStyle"), qn("w:tblStyle")}
    for node in element.iter():
        if node.tag not in style_tags:
            continue
        old_style_id = node.get(qn("w:val"))
        if old_style_id in style_id_map:
            node.set(qn("w:val"), style_id_map[old_style_id])


def append_docx_body_preserve_relationships(
    target_doc: Document,
    source_doc: Document,
    *,
    remap_styles_by_name: bool = True,
) -> None:
    """Append source body XML to target, remapping relationships and styles."""
    target_body = target_doc.element.body
    sect_pr = target_body.sectPr
    if sect_pr is not None:
        target_body.remove(sect_pr)

    relmap: dict[str, str] = {}
    rel_attrs = {qn("r:id"), qn("r:embed"), qn("r:link")}
    style_id_map = build_style_id_map(target_doc, source_doc) if remap_styles_by_name else {}

    def remap_relationships(element) -> None:
        for node in element.iter():
            for attr_name in list(node.attrib):
                if attr_name not in rel_attrs:
                    continue
                old_rid = node.attrib[attr_name]
                if old_rid not in source_doc.part.rels:
                    continue
                if old_rid not in relmap:
                    rel = source_doc.part.rels[old_rid]
                    if rel.is_external:
                        relmap[old_rid] = target_doc.part.relate_to(
                            rel.target_ref, rel.reltype, is_external=True
                        )
                    else:
                        relmap[old_rid] = target_doc.part.relate_to(
                            rel.target_part, rel.reltype
                        )
                node.attrib[attr_name] = relmap[old_rid]

    for child in list(source_doc.element.body):
        if child.tag == qn("w:sectPr"):
            continue
        copied = copy.deepcopy(child)
        remap_relationships(copied)
        remap_style_ids(copied, style_id_map)
        target_body.append(copied)

    if sect_pr is not None:
        target_body.append(sect_pr)


def clear_section_references(
    doc: Document,
    *,
    headers: bool = False,
    footers: bool = False,
) -> None:
    """Remove inherited section header/footer references.

    Use this after deleting template sample sections if generated body pages
    inherit a stale back-matter header such as 致谢 or 参考文献.
    """
    for section in doc.sections:
        sect_pr = section._sectPr
        if headers:
            for ref in list(sect_pr.findall(qn("w:headerReference"))):
                sect_pr.remove(ref)
        if footers:
            for ref in list(sect_pr.findall(qn("w:footerReference"))):
                sect_pr.remove(ref)


def force_heading_page_breaks(
    doc: Document, heading_styles: tuple[str, ...] = DEFAULT_UNNUMBERED_HEADING_STYLES
) -> None:
    for paragraph in doc.paragraphs:
        if paragraph.style.name in heading_styles:
            paragraph.paragraph_format.page_break_before = True


def set_cell_border(cell, edge: str, *, val: str = "nil", size: int = 0) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = tc_pr.find(qn("w:tcBorders"))
    if tc_borders is None:
        tc_borders = OxmlElement("w:tcBorders")
        tc_pr.append(tc_borders)
    border = tc_borders.find(qn(f"w:{edge}"))
    if border is None:
        border = OxmlElement(f"w:{edge}")
        tc_borders.append(border)
    border.set(qn("w:val"), val)
    if val != "nil":
        border.set(qn("w:sz"), str(size))
        border.set(qn("w:space"), "0")
        border.set(qn("w:color"), "000000")


def format_three_line_tables(
    doc: Document,
    *,
    font_name: str | None = None,
    east_asia_font: str | None = None,
    font_size_pt: float | None = None,
    start_marker: str | None = None,
    end_marker: str | None = None,
    include_start_marker: bool = True,
) -> None:
    """Coerce every table into a three-line table.

    This is opinionated and meant for Chinese-thesis-style three-line tables.
    Don't enable it unless your template actually requires this layout —
    activate via --three-line-tables or enable_three_line_tables in config.
    """
    allowed_tables = table_ids_in_marker_scope(
        doc,
        start_marker=start_marker,
        end_marker=end_marker,
        include_start_marker=include_start_marker,
    )
    for table in doc.tables:
        if allowed_tables is not None and id(table._tbl) not in allowed_tables:
            continue
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        for row in table.rows:
            for cell in row.cells:
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
                    set_cell_border(cell, edge, val="nil")
                for paragraph in cell.paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    clear_paragraph_indent(paragraph)
                    for run in paragraph.runs:
                        if font_size_pt is not None:
                            run.font.size = Pt(font_size_pt)
                        if font_name is not None:
                            run.font.name = font_name
                        if east_asia_font is not None:
                            set_east_asia_font(run, east_asia_font)
        if table.rows:
            for cell in table.rows[0].cells:
                set_cell_border(cell, "top", val="single", size=12)
                set_cell_border(cell, "bottom", val="single", size=8)
            for cell in table.rows[-1].cells:
                set_cell_border(cell, "bottom", val="single", size=12)


def normalize_hyperlinks_black(
    doc: Document,
    *,
    start_marker: str | None = None,
    end_marker: str | None = None,
    include_start_marker: bool = True,
) -> None:
    elements = list(
        body_elements_in_marker_scope(
            doc,
            start_marker=start_marker,
            end_marker=end_marker,
            include_start_marker=include_start_marker,
        )
    )
    for hyperlink in (link for element in elements for link in element.iter(qn("w:hyperlink"))):
        for run_el in hyperlink.findall(qn("w:r")):
            rpr = run_el.find(qn("w:rPr"))
            if rpr is None:
                rpr = OxmlElement("w:rPr")
                run_el.insert(0, rpr)
            color = rpr.find(qn("w:color"))
            if color is None:
                color = OxmlElement("w:color")
                rpr.append(color)
            color.set(qn("w:val"), "000000")
            u = rpr.find(qn("w:u"))
            if u is None:
                u = OxmlElement("w:u")
                rpr.append(u)
            u.set(qn("w:val"), "none")


def apply_basic_style_mapping(
    doc: Document,
    *,
    body_style: str,
    unnumbered_h1: set[str],
    unnumbered_heading_styles: tuple[str, ...],
    caption_regex: str,
    body_candidate_styles: tuple[str, ...],
    body_font_name: str | None,
    body_east_asia_font: str | None,
    body_font_size_pt: float | None,
    start_marker: str | None = None,
    end_marker: str | None = None,
    include_start_marker: bool = True,
) -> None:
    caption_re = re.compile(caption_regex)
    for paragraph in iter_paragraphs_in_marker_scope(
        doc,
        start_marker=start_marker,
        end_marker=end_marker,
        include_start_marker=include_start_marker,
    ):
        text = re.sub(r"\s+", "", paragraph.text)
        if (
            paragraph.style.name in unnumbered_heading_styles
            and text in unnumbered_h1
        ):
            suppress_heading_number(paragraph)
            continue
        if "<w:drawing" in paragraph._element.xml:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            continue
        if caption_re.match(paragraph.text.strip()):
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            paragraph.paragraph_format.first_line_indent = None
            paragraph.paragraph_format.space_before = Pt(6)
            paragraph.paragraph_format.space_after = Pt(6)
            for run in paragraph.runs:
                run.bold = True
                if body_font_size_pt is not None:
                    run.font.size = Pt(body_font_size_pt)
            continue
        if paragraph.style.name in body_candidate_styles and paragraph.text.strip():
            if body_style and body_style in doc.styles:
                try:
                    paragraph.style = doc.styles[body_style]
                except KeyError:
                    # body_style 在某些 python-docx 版本里若用 style_id 索引会 KeyError
                    pass
            for run in paragraph.runs:
                if body_font_name is not None:
                    run.font.name = body_font_name
                if body_font_size_pt is not None:
                    run.font.size = Pt(body_font_size_pt)
                if body_east_asia_font is not None:
                    set_east_asia_font(run, body_east_asia_font)


def load_config(path: Path | None) -> dict:
    if path is None:
        return {}
    return json.loads(path.read_text(encoding="utf-8"))


def main() -> int:
    parser = argparse.ArgumentParser()
    parser.add_argument("--template", required=True)
    parser.add_argument("--body-docx", default=None)
    parser.add_argument("--out", required=True)
    parser.add_argument("--source-md", default=None)
    parser.add_argument("--config", default=None, help="JSON config with style names/rules")
    parser.add_argument(
        "--three-line-tables",
        action="store_true",
        help="Coerce all tables into Chinese-thesis-style three-line tables (opt-in)",
    )
    parser.add_argument(
        "--keep-hyperlink-color",
        action="store_true",
        help="Skip forcing hyperlinks to black/no underline",
    )
    args = parser.parse_args()

    cfg = load_config(Path(args.config) if args.config else None)
    template_doc = Document(args.template)
    body_doc = Document(args.body_docx) if args.body_docx else None

    if args.source_md:
        rebuild_from_markdown(
            template_doc,
            Path(args.source_md).read_text(encoding="utf-8"),
            table_captions=cfg.get("table_captions", []),
            body_east_asia_font=cfg.get("body_east_asia_font"),
            insert_toc=cfg.get("insert_toc", True),
        )
    else:
        if body_doc is None:
            parser.error("--body-docx is required when --source-md is not provided")
        append_docx_body_preserve_relationships(
            template_doc,
            body_doc,
            remap_styles_by_name=cfg.get("remap_styles_by_name", True),
        )
    formatting_start_marker = cfg.get("formatting_start_marker")
    formatting_end_marker = cfg.get("formatting_end_marker")
    formatting_include_start_marker = cfg.get("formatting_include_start_marker", True)

    apply_basic_style_mapping(
        template_doc,
        body_style=cfg.get("body_style", "Normal"),
        unnumbered_h1=set(cfg.get("unnumbered_h1", [])),
        unnumbered_heading_styles=tuple(
            cfg.get("unnumbered_heading_styles", DEFAULT_UNNUMBERED_HEADING_STYLES)
        ),
        caption_regex=cfg.get("caption_regex", DEFAULT_CAPTION_REGEX),
        body_candidate_styles=tuple(
            cfg.get("body_candidate_styles", DEFAULT_BODY_CANDIDATE_STYLES)
        ),
        body_font_name=cfg.get("body_font_name"),
        body_east_asia_font=cfg.get("body_east_asia_font"),
        body_font_size_pt=cfg.get("body_font_size_pt"),
        start_marker=formatting_start_marker,
        end_marker=formatting_end_marker,
        include_start_marker=formatting_include_start_marker,
    )

    enable_three_line_tables = args.three_line_tables or cfg.get(
        "enable_three_line_tables", False
    )
    if enable_three_line_tables:
        format_three_line_tables(
            template_doc,
            font_name=cfg.get("table_font_name"),
            east_asia_font=cfg.get("table_east_asia_font"),
            font_size_pt=cfg.get("table_font_size_pt"),
            start_marker=formatting_start_marker,
            end_marker=formatting_end_marker,
            include_start_marker=formatting_include_start_marker,
        )

    enable_black_hyperlinks = (not args.keep_hyperlink_color) and cfg.get(
        "enable_black_hyperlinks", True
    )
    if enable_black_hyperlinks:
        normalize_hyperlinks_black(
            template_doc,
            start_marker=formatting_start_marker,
            end_marker=formatting_end_marker,
            include_start_marker=formatting_include_start_marker,
        )

    if cfg.get("heading_1_page_breaks", False):
        force_heading_page_breaks(
            template_doc,
            tuple(cfg.get("heading_1_styles", DEFAULT_UNNUMBERED_HEADING_STYLES)),
        )

    clear_section_references(
        template_doc,
        headers=cfg.get("clear_header_references", False),
        footers=cfg.get("clear_footer_references", False),
    )

    out = Path(args.out)
    out.parent.mkdir(parents=True, exist_ok=True)
    template_doc.save(out)
    print(out)
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
